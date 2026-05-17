from __future__ import annotations

"""Document parsers: PDF, DOCX, XLSX, TXT/MD."""

import logging
import warnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedChunk:
    """Raw text block extracted from a source document."""
    text: str
    source: str          # original filename (basename)
    page: Optional[int] = None
    section: Optional[str] = None


# ---------------------------------------------------------------------------
# Routing
# ---------------------------------------------------------------------------

def parse_file(path: Path) -> list[ParsedChunk]:
    """Dispatch to the right parser based on file suffix. Silently skips on error."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _parse_pdf(path)
        elif suffix == ".docx":
            return _parse_docx(path)
        elif suffix in (".xlsx", ".xls"):
            return _parse_xlsx(path)
        elif suffix in (".txt", ".md"):
            return _parse_text(path)
        else:
            logger.warning("Unsupported file type: %s — skipping", path.name)
            return []
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to parse %s: %s — skipping", path.name, exc)
        return []


def walk_corpus(corpus_dir: Path) -> list[ParsedChunk]:
    """Recursively parse all supported docs in corpus_dir."""
    chunks: list[ParsedChunk] = []
    supported = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md"}
    for p in sorted(corpus_dir.rglob("*")):
        if p.is_file() and p.suffix.lower() in supported:
            logger.info("Parsing %s", p.name)
            chunks.extend(parse_file(p))
    logger.info("Parsed %d raw chunks from %s", len(chunks), corpus_dir)
    return chunks


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> list[ParsedChunk]:
    from pypdf import PdfReader  # type: ignore

    source = path.name
    reader = PdfReader(str(path))
    chunks: list[ParsedChunk] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            chunks.append(ParsedChunk(text=text, source=source, page=page_num))
    return chunks


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> list[ParsedChunk]:
    import docx  # type: ignore

    source = path.name
    doc = docx.Document(str(path))
    chunks: list[ParsedChunk] = []
    current_section: Optional[str] = None
    buffer: list[str] = []

    def flush(section: Optional[str]) -> None:
        text = "\n".join(buffer).strip()
        if text:
            chunks.append(ParsedChunk(text=text, source=source, section=section))
        buffer.clear()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            flush(current_section)
            current_section = para.text.strip() or current_section
        else:
            t = para.text.strip()
            if t:
                buffer.append(t)

    # Tables
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(row_texts))
        table_block = "\n".join(rows_text).strip()
        if table_block:
            chunks.append(ParsedChunk(text=table_block, source=source, section=current_section))

    flush(current_section)
    return chunks


# ---------------------------------------------------------------------------
# XLSX (source knowledge base files, not the input RFP)
# ---------------------------------------------------------------------------

def _parse_xlsx(path: Path) -> list[ParsedChunk]:
    import openpyxl  # type: ignore

    source = path.name
    chunks: list[ParsedChunk] = []
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            row_str = " | ".join(cells).strip(" |")
            if row_str:
                rows_text.append(row_str)
        text = "\n".join(rows_text).strip()
        if text:
            chunks.append(ParsedChunk(text=text, source=source, section=sheet_name))
    wb.close()
    return chunks


# ---------------------------------------------------------------------------
# TXT / MD
# ---------------------------------------------------------------------------

def _parse_text(path: Path) -> list[ParsedChunk]:
    source = path.name
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if text:
        return [ParsedChunk(text=text, source=source)]
    return []
